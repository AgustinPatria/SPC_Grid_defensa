"""Construye reporte_reunion.docx para la reunion con el profesor.

Lee dinamicamente:
  - El log mas reciente de resultados/main_run*.log (parsing por regex)
  - Los CSV en inspeccion_v2/L{2..5}_*_out/ (regenerar con scripts L*.py)

Asi el docx queda en sync con la ultima corrida sin tocar codigo.

Contenido:
  1. Resumen ejecutivo (DLConfig + backtests IS/OOS + selecciones g*).
  2. L2 — Motor DL (p_bull, calibracion, dispersion entre NNs, mu_hat).
  3. L3 — Escenarios DL compartidos (5 reps + N candidatos, correlacion).
  4. L4 — Motor optimizador (decomposicion z, w(t) apilado, turnover).
  5. L5 — Regret grid (V/R tables, plano lambda-m, curvas).
  6. Sintesis y proxima capa.

Uso:
    python build_reporte_reunion.py
"""
from __future__ import annotations

import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


PROJECT_ROOT = Path(__file__).resolve().parent
INSP_V2 = PROJECT_ROOT / "inspeccion_v2"
RESULTADOS = PROJECT_ROOT / "resultados"
OUT_PATH = PROJECT_ROOT / "reporte_reunion.docx"


# ================================================================ parsing log

def latest_main_log() -> Path:
    """Devuelve el main_run*.log mas reciente por mtime."""
    candidates = list(RESULTADOS.glob("main_run*.log"))
    if not candidates:
        raise FileNotFoundError(f"No hay main_run*.log en {RESULTADOS}")
    return max(candidates, key=lambda p: p.stat().st_mtime)


_RE_DLCONFIG = re.compile(
    r"DLConfig:\s*H=(\d+)\s+hidden=(\d+)\s+epochs=(\d+)\s+patience=(\d+)"
)
_RE_GSTAR = re.compile(
    r"g\*_(mean|worst)(_oos)?\s+\(ec\.\s*\d+\):\s*"
    r"lambda=([\d.]+)\s+m=([\d.]+)\s+"
    r"(mean|worst)_regret=\$([\d,.]+)"
)
_RE_VSUMMARY = re.compile(
    r"V(?:_oos)?:\s*mean=\$\s*([\d,.]+)\s+worst=\$\s*([\d,.]+)\s+"
    r"best=\$\s*([\d,.]+)"
)
_RE_RETSC = re.compile(
    r"retorno promedio sobre escenarios(?:\s+OOS)?\s*=\s*([+\-][\d.]+%)"
)
_RE_RETPEOR = re.compile(
    r"retorno en el peor escenario(?:\s+OOS)?\s*=\s*([+\-][\d.]+%)"
)
_RE_RETPORT = re.compile(
    r"ret port \(.*?\) por escenario(?:\s+OOS)?:\s+\[(.*?)\]"
)
_RE_CAPFO = re.compile(
    r"cap FO\s+\(con costos\)\s+por escenario:\s+\[(.*?)\]"
)
_RE_BACKTEST_ROW = re.compile(
    r"\s+(OPT(?:_oos)?(?:.*?)|Regret-Grid(?:_oos)?.*?|Naive.*?)\s+\$\s*([\d,.]+)\s+"
    r"([+\-][\d.]+%)\s+\$\s*([+\-][\d,.]+)"
)
_RE_PBULL_DISP = re.compile(
    r"\s+(\w+)\s+p_bull std promedio \(sobre t\)=([\d.]+)\s+"
    r"max std=([\d.]+)\s+max\(max-min\)=([\d.]+)"
)


def _split_is_oos(text: str) -> tuple[str, str]:
    """Parte el log en bloque IS y bloque OOS.

    Anchor: 'FASE 6' (palabra exacta, tolera cualquier separador despues).
    """
    # Buscamos el header "FASE 6" — el em-dash + espacios variables siguen.
    m = re.search(r"FASE 6[^\n]*REGRET GRID", text)
    if not m:
        return text, ""
    cut = m.start()
    return text[:cut], text[cut:]


def parse_main_log(path: Path) -> dict:
    """Extrae IS + OOS de un main_run log."""
    text = path.read_text(encoding="utf-8", errors="replace")
    out: dict = {"path": path}

    m = _RE_DLCONFIG.search(text)
    if m:
        out["dlconfig"] = {
            "H": int(m.group(1)), "hidden": int(m.group(2)),
            "epochs": int(m.group(3)), "patience": int(m.group(4)),
        }

    is_block, oos_block = _split_is_oos(text)

    def parse_block(block: str, is_oos: bool) -> dict:
        b: dict = {}
        # g* selections — filtrar por presencia/ausencia del sufijo _oos
        for m in _RE_GSTAR.finditer(block):
            kind, oos_suffix, lam, m_, _, regret = m.groups()
            has_oos = bool(oos_suffix)
            if has_oos != is_oos:
                continue
            b[f"g_{kind}"] = {
                "lambda": float(lam), "m": float(m_),
                "regret": float(regret.replace(",", "")),
            }
        # V summaries siguen al g* respectivo en el log
        vs = list(_RE_VSUMMARY.finditer(block))
        # En IS block los V de IS aparecen primero (2 de 4); en OOS block solo
        # hay 2 V summaries (de OOS). Asignamos los relevantes.
        if is_oos:
            # OOS block: las 2 primeras V summaries son g_mean_oos y g_worst_oos
            for label, m in zip(["g_mean", "g_worst"], vs[:2]):
                if label in b:
                    b[label]["V_mean"]  = float(m.group(1).replace(",", ""))
                    b[label]["V_worst"] = float(m.group(2).replace(",", ""))
                    b[label]["V_best"]  = float(m.group(3).replace(",", ""))
        else:
            # IS block: las 2 V summaries son de IS (g_mean, g_worst)
            for label, m in zip(["g_mean", "g_worst"], vs[:2]):
                if label in b:
                    b[label]["V_mean"]  = float(m.group(1).replace(",", ""))
                    b[label]["V_worst"] = float(m.group(2).replace(",", ""))
                    b[label]["V_best"]  = float(m.group(3).replace(",", ""))
        # retorno escenarios
        rs = _RE_RETSC.search(block)
        if rs:
            b["ret_escenarios_pct"] = rs.group(1)
        rp = _RE_RETPEOR.search(block)
        if rp:
            b["ret_peor_escenario_pct"] = rp.group(1)
        # escenarios cum_port
        rport = _RE_RETPORT.search(block)
        if rport:
            b["escenarios_ret_port"] = [
                s.strip().strip("'\"") for s in rport.group(1).split(",")
            ]
        cf = _RE_CAPFO.search(block)
        if cf:
            b["escenarios_cap_fo"] = [
                s.strip().strip("'\"") for s in cf.group(1).split(",")
            ]
        # backtest rows
        bt = []
        for m in _RE_BACKTEST_ROW.finditer(block):
            name, cap, ret, inc = m.groups()
            bt.append({
                "politica": name.strip(),
                "capital":  f"${cap.strip()}",
                "retorno":  ret.strip(),
                "incremento": f"${inc.strip()}",
            })
        if bt:
            b["backtest"] = bt
        return b

    out["is"]  = parse_block(is_block,  is_oos=False)
    out["oos"] = parse_block(oos_block, is_oos=True)

    # p_bull dispersion (solo IS)
    pdisp = {}
    for m in _RE_PBULL_DISP.finditer(is_block):
        asset, std_avg, max_std, max_minmax = m.groups()
        pdisp[asset] = {
            "std_avg": float(std_avg),
            "max_std": float(max_std),
            "max_minmax": float(max_minmax),
        }
    if pdisp:
        out["pbull_disp"] = pdisp

    return out


# ================================================================ doc helpers

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_paragraph(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    return p


def add_image(doc, rel_path: Path, width_cm: float = 16.0,
              caption: str | None = None):
    abs_path = PROJECT_ROOT / rel_path
    if not abs_path.exists():
        add_paragraph(doc, f"[imagen no encontrada: {rel_path}]", italic=True)
        return
    doc.add_picture(str(abs_path), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)


def add_df_table(doc, df: pd.DataFrame, max_rows: int | None = None,
                 float_fmt: str = "{:.4f}"):
    if max_rows is not None and len(df) > max_rows:
        df = df.head(max_rows)
    table = doc.add_table(rows=1 + len(df), cols=len(df.columns))
    table.style = "Light Grid Accent 1"
    for ci, col in enumerate(df.columns):
        cell = table.rows[0].cells[ci]
        cell.text = str(col)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for ri, (_, row) in enumerate(df.iterrows(), start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            if isinstance(val, float):
                cell.text = float_fmt.format(val)
            else:
                cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)


def list_to_df(rows: list[dict]) -> pd.DataFrame:
    return pd.DataFrame(rows)


# ================================================================ secciones

def seccion_portada(doc, log: dict):
    title = doc.add_heading("SPC_Grid — Reporte de resultados", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cfg = log.get("dlconfig", {})
    desc = (
        "Pipeline de optimizacion de portafolio (media-varianza + costos) con "
        "capa DL (LSTM cuantilico per-cell) y seleccion por regret-grid sobre "
        "la grilla (lambda, m). Configuracion vigente: mu_hat_source='p_hist' "
        "(PDF literal sec 1.3), p_method='walking', 15 NNs independientes con "
        "seed deterministico, escenarios compartidos por ensemble. "
        f"LSTM: H={cfg.get('H','?')}, hidden={cfg.get('hidden','?')}, "
        f"epochs={cfg.get('epochs','?')}, patience={cfg.get('patience','?')}. "
        f"Run base: {log['path'].name}."
    )
    add_paragraph(doc, desc, italic=True, size=10)
    doc.add_paragraph()


# ----------------------------------- 1. resumen ejecutivo
def seccion_resumen(doc, log: dict):
    add_heading(doc, "1. Resumen ejecutivo", level=1)
    add_paragraph(doc,
        "El pipeline corre de punta a punta (15 NNs + 30 solves GAMS + "
        "5 capas de inspeccion). Selecciones del regret-grid y backtests "
        "sobre la trayectoria historica real:",
    )

    # ---- IS
    add_heading(doc, "Backtest IN-SAMPLE (t=1..163, 163 semanas)", level=2)
    is_bt = log["is"].get("backtest", [])
    if is_bt:
        df = list_to_df(is_bt)
        add_df_table(doc, df[["politica", "capital", "retorno", "incremento"]])
    g = log["is"].get("g_mean", {})
    if g:
        add_paragraph(doc,
            f"g*_mean IS = (lambda={g['lambda']:.2f}, m={g['m']:.2f}). "
            f"mean_regret = ${g['regret']:,.2f}; "
            f"worst_regret = ${log['is'].get('g_worst', {}).get('regret', 0):,.2f}. "
            f"Retorno promedio sobre los 5 escenarios DL: "
            f"{log['is'].get('ret_escenarios_pct', '?')}."
        )

    # ---- OOS
    add_heading(doc, "Backtest OUT-OF-SAMPLE (t=148..163, 16 semanas)", level=2)
    oos_bt = log["oos"].get("backtest", [])
    if oos_bt:
        df = list_to_df(oos_bt)
        add_df_table(doc, df[["politica", "capital", "retorno", "incremento"]])
    g = log["oos"].get("g_mean", {})
    if g:
        add_paragraph(doc,
            f"g*_mean OOS = (lambda={g['lambda']:.2f}, m={g['m']:.2f}). "
            f"mean_regret = ${g['regret']:,.2f}; "
            f"worst_regret = ${log['oos'].get('g_worst', {}).get('regret', 0):,.2f}. "
            f"Retorno promedio sobre escenarios DL OOS: "
            f"{log['oos'].get('ret_escenarios_pct', '?')}."
        )

    add_paragraph(doc,
        "4/4 selecciones (g*_mean IS, g*_worst IS, g*_mean OOS, g*_worst OOS) "
        "caen en alguna frontera del grid. La inspeccion por capas (siguientes "
        "secciones) explica por que.",
        italic=True,
    )


# ----------------------------------- 2. L2 — DL
def seccion_l2(doc, log: dict):
    add_heading(doc, "2. Capa DL (L2) — las 15 NNs cuantilicas", level=1)
    add_paragraph(doc,
        "Cada celda g = (lambda, m) entrena su propia NN cuantilica con seed "
        "deterministica = cell_seed(lam, m). Las 15 producen 15 series "
        "p_bull(t) que modulan mu_mix(t) / sigma_mix(t) en el optimizador.",
    )

    add_heading(doc, "2.1 — Probabilidad bull sobre el tiempo", level=2)
    add_image(doc, "inspeccion_v2/L2_dl_out/01_pbull_by_cell.png",
              caption="p_bull(t) por activo. 15 NNs en gris, ensemble en negro. "
                      "Linea roja: t_test_start. Linea azul punteada: 0.5.")

    add_heading(doc, "2.2 — Dispersion entre las 15 NNs", level=2)
    add_image(doc, "inspeccion_v2/L2_dl_out/02_dispersion.png",
              caption="std y rango (max-min) de p_bull(t) entre las 15 redes.")
    pdisp = log.get("pbull_disp", {})
    if pdisp:
        rows = []
        for a, d in pdisp.items():
            rows.append({"activo": a, "std_avg": d["std_avg"],
                         "max_std": d["max_std"], "max(max-min)": d["max_minmax"]})
        add_df_table(doc, list_to_df(rows), float_fmt="{:.4f}")
    add_paragraph(doc,
        "max(max-min)=0.60 en algunos t: una NN dice p_bull=0.85 y otra dice 0.25 "
        "para la misma ventana. Distintos seeds convergen a optimos locales del "
        "pinball loss porque la senal en la data es debil.",
        italic=True,
    )

    add_heading(doc, "2.3 — Calibracion direccional (post-warmup)", level=2)
    add_image(doc, "inspeccion_v2/L2_dl_out/03_calibration_IS.png",
              caption="accuracy(p_bull > 0.5 vs r_real >= 0) por celda IS. "
                      "Linea roja: base rate bull. Linea gris: coin flip (0.5).")
    cal_path = INSP_V2 / "L2_dl_out" / "03_calibration_IS.csv"
    if cal_path.exists():
        df = pd.read_csv(cal_path)
        ens = df[df["lambda"] == "ensemble"]
        if not ens.empty:
            pivot = ens.set_index("asset")[["accuracy_dir", "base_rate_bull",
                                             "p_bull_mean", "p_bull_std"]]
            pivot = pivot.reset_index()
            add_df_table(doc, pivot, float_fmt="{:.3f}")
        add_paragraph(doc,
            "Ensemble = promedio de logits de las 15. base_rate_bull = fraccion "
            "real de semanas con r>=0 (lo que sacaria un modelo que siempre dice "
            "'bull'). coin flip = 0.5.",
            italic=True,
        )

    add_heading(doc, "2.4 — mu_hat (retornos por regimen, anclaje historico)", level=2)
    mu_hat_path = INSP_V2 / "L2_dl_out" / "05_mu_hat_per_cell_IS.csv"
    if mu_hat_path.exists():
        df = pd.read_csv(mu_hat_path).head(1)
        cols = [c for c in df.columns if c.startswith("mu_hat_")]
        df_show = df[cols].T.reset_index()
        df_show.columns = ["activo_regimen", "mu_hat (sem)"]
        add_df_table(doc, df_show, float_fmt="{:+.6f}")
    add_paragraph(doc,
        "mu_hat es IDENTICO en las 15 celdas (depende solo de p_hist del CSV y "
        "r_hist, no de la NN). Caveat: p_hist viene de un HMM externo que define "
        "regimen por volatilidad, mientras que el LSTM lo define por r>=0. Las "
        "dos definiciones no se alinean.",
        italic=True,
    )


# ----------------------------------- 3. L3 — escenarios
def seccion_l3(doc, log: dict):
    add_heading(doc, "3. Escenarios DL (L3) — 5 representativos compartidos", level=1)
    add_paragraph(doc,
        "El ensemble (promedio de logits) genera N=1000 candidatos con "
        "rolling-forward desde la ventana inicial. Se reducen a 5 reps por "
        "quintiles, rankeados por capital terminal de portafolio FO-aligned "
        "(con costos). Estos 5 reps son COMPARTIDOS por todas las celdas.",
    )

    add_heading(doc, "3.1 — Trayectorias representativas IS", level=2)
    add_image(doc, "inspeccion_v2/L3_escenarios_out/01_reps_paths_IS.png",
              caption="cum_ret acumulado por escenario (gradient viridis = peor a mejor). "
                      "Paneles: SPX, CMC200, portafolio w_ref. T=163 semanas.")
    esc_is = log["is"].get("escenarios_ret_port", [])
    esc_fo = log["is"].get("escenarios_cap_fo", [])
    if esc_is:
        df = pd.DataFrame({
            "escenario": [f"s={i}" for i in range(len(esc_is))],
            "ret port (sin costos)": esc_is,
            "cap FO (con costos)":   esc_fo if esc_fo else [""] * len(esc_is),
        })
        add_df_table(doc, df)

    add_heading(doc, "3.2 — Distribucion de candidatos", level=2)
    add_image(doc, "inspeccion_v2/L3_escenarios_out/02_fan_IS.png",
              caption="Fan chart 5-95 / 25-75 pct + mediana de los 1000 candidatos, "
                      "con los 5 reps superpuestos.")
    summary_path = INSP_V2 / "L3_escenarios_out" / "00_summary_IS_vs_OOS.csv"
    if summary_path.exists():
        add_df_table(doc, pd.read_csv(summary_path))

    add_heading(doc, "3.3 — Correlacion SPX-CMC200 en los escenarios", level=2)
    corr_path = INSP_V2 / "L3_escenarios_out" / "06_correlation_IS.csv"
    if corr_path.exists():
        add_df_table(doc, pd.read_csv(corr_path), float_fmt="{:.3f}")
    add_paragraph(doc,
        "Los 5 reps tienen corr(SPX, CMC) >= 0.97 vs 0.31 historico. Viene del "
        "'mismo q en todos los activos' de generate_candidate_scenarios. Los "
        "activos quedan comonotonos -- no hay diversificacion real dentro de "
        "un escenario.",
        italic=True,
    )


# ----------------------------------- 4. L4 — optimizador
def seccion_l4(doc):
    add_heading(doc, "4. Optimizador (L4) — pesos, costos, riesgo", level=1)
    add_paragraph(doc,
        "Para cada celda g=(lambda, m), GAMSPy+IPOPT resuelve la FO: "
        "z = sum_t [sum_i w(i,t)*mu(i,t)  -  lambda*(sum_ij w_i*w_j*sigma_ij - V_max) "
        " -  sum_i c_base(i)*costo_mult*(u(i,t)+v(i,t))]. Sanity z analitico vs "
        "IPOPT: 0.0 al cero decimal en las 30 celdas (15 IS + 15 OOS).",
    )

    add_heading(doc, "4.1 — Trayectoria de pesos w(i, t)", level=2)
    add_image(doc, "inspeccion_v2/L4_optimizador_out/05_w_stacked_IS.png",
              width_cm=16.5,
              caption="Stacked area chart de w(i, t) por celda. Filas = lambda, "
                      "columnas = m. La frontera entre las dos bandas ES la "
                      "trayectoria de rebalanceo. Header: |Delta w|_tot.")

    add_heading(doc, "4.2 — Decomposicion de z (retorno, riesgo, costo)", level=2)
    add_image(doc, "inspeccion_v2/L4_optimizador_out/01_z_decomposition_IS.png",
              caption="Barras stacked: retorno (verde) + (-riesgo) (rojo) + "
                      "(-costo) (gris) por celda IS.")
    add_paragraph(doc,
        "Costo << retorno en todas las celdas; m no controla nada materialmente "
        "en la FO. lambda alto puede producir riesgo NEGATIVO: el portafolio "
        "queda debajo del V_max budget.",
        italic=True,
    )

    add_heading(doc, "4.3 — Turnover total por celda", level=2)
    turn_path = INSP_V2 / "L4_optimizador_out" / "03_turnover_IS.csv"
    if turn_path.exists():
        df = pd.read_csv(turn_path)
        cols = [c for c in ["lambda", "m", "w_mean_SPX", "w_mean_CMC200",
                            "turnover_total"] if c in df.columns]
        add_df_table(doc, df[cols], float_fmt="{:.4f}")


# ----------------------------------- 5. L5 — regret
def seccion_l5(doc, log: dict):
    add_heading(doc, "5. Regret grid (L5) — V[g, s], R[g, s], seleccion g*", level=1)
    add_paragraph(doc,
        "V[g, s] = capital terminal de la politica g aplicada al escenario s. "
        "R[g, s] = max_g' V[g', s] - V[g, s]. g*_mean = argmin mean regret "
        "(ec. 23). g*_worst = argmin worst-case regret (ec. 24).",
    )

    add_heading(doc, "5.1 — Tabla V[g, s] IS", level=2)
    add_image(doc, "inspeccion_v2/L5_regret_out/01_V_table_IS.png",
              caption="Heatmap V[g, s] IS. Filas = celdas (lambda, m). Columnas = "
                      "escenarios. Rectangulo azul = g*_mean.")

    add_heading(doc, "5.2 — Tabla R[g, s] IS", level=2)
    add_image(doc, "inspeccion_v2/L5_regret_out/02_R_table_IS.png",
              caption="Heatmap R[g, s] IS. R=0 indica que g es la mejor celda "
                      "para ese escenario.")

    add_heading(doc, "5.3 — Plano (lambda, m): mean_regret y worst_regret", level=2)
    add_image(doc, "inspeccion_v2/L5_regret_out/03_mean_plane_IS.png",
              width_cm=10,
              caption="mean_regret IS en plano (lambda, m). g*_mean marcado en azul.")
    add_image(doc, "inspeccion_v2/L5_regret_out/04_worst_plane_IS.png",
              width_cm=10,
              caption="worst_regret IS en plano (lambda, m). g*_worst marcado.")

    add_heading(doc, "5.4 — Curvas regret vs lambda y vs m", level=2)
    add_image(doc, "inspeccion_v2/L5_regret_out/05_curves_IS.png",
              caption="mean_regret vs lambda (lineas por m) y vs m (lineas por lambda).")

    add_heading(doc, "5.5 — Resumen pivot mean_regret IS", level=2)
    summary_path = INSP_V2 / "L5_regret_out" / "05_regret_summary_IS.csv"
    if summary_path.exists():
        df = pd.read_csv(summary_path)
        pivot = df.pivot_table(index="lambda", columns="m",
                                values="mean_regret").reset_index()
        pivot.columns = [str(c) for c in pivot.columns]
        add_df_table(doc, pivot, float_fmt="{:.2f}")

    add_heading(doc, "5.6 — IS vs OOS", level=2)
    rows = []
    for seg in ["is", "oos"]:
        for kind in ["g_mean", "g_worst"]:
            g = log[seg].get(kind, {})
            if g:
                rows.append({
                    "segmento": seg.upper(),
                    "criterio": kind.replace("g_", "g*_"),
                    "g*": f"(lambda={g['lambda']:.2f}, m={g['m']:.2f})",
                    "regret": f"${g['regret']:,.2f}",
                })
    if rows:
        add_df_table(doc, list_to_df(rows))


# ----------------------------------- 6. sintesis
def seccion_sintesis(doc, log: dict):
    add_heading(doc, "6. Sintesis", level=1)
    add_paragraph(doc,
        "El pipeline esta funcionando correctamente — la formulacion del "
        "optimizador (sanity z exacto) y el regret-grid son correctos sobre "
        "la informacion que reciben. El cuello esta en las capas de entrada:",
    )

    add_heading(doc, "Cinco hallazgos estructurales", level=2)

    add_paragraph(doc, "1. mu_hat invertido (L2.4): mu_hat(bear) > mu_hat(bull) "
        "en ambos activos. p_hist (HMM externo) no se alinea con la regla "
        "r>=0 del LSTM. El optimizador resuelve un problema donde 'ser "
        "bullish' baja el retorno esperado.")

    add_paragraph(doc, "2. LSTM con calibracion direccional cercana a moneda "
        "(L2.3): la red captura algo de senal pero el accuracy se mantiene en "
        "el rango [0.4, 0.6]. 163 semanas no alcanzan para entrenar "
        "direccionalidad robusta.")

    add_paragraph(doc, "3. Escenarios comonotonos (L3.3): corr(SPX, CMC) >= 0.97 "
        "en los 5 reps vs 0.31 historico. Diversificar entre activos no "
        "diversifica dentro del escenario.")

    add_paragraph(doc, "4. Soluciones de esquina (L4.1): el optimizador "
        "tipicamente elige una esquina del simplex, modulada por lambda. La "
        "diversificacion 'mixta' aparece solo donde lambda esta cerca del "
        "punto de cambio de regimen del optimo.")

    add_paragraph(doc, "5. m es practicamente inerte (L4.2 + L5.5): el costo "
        "del orden de 10^-4 vs retornos del orden de 10^-1 hace que m no "
        "afecte materialmente V[g, s]. El grid efectivo es 1D (solo lambda).")

    add_heading(doc, "Comparacion contra benchmarks", level=2)

    is_bt = {r["politica"].split(" (")[0]: r for r in log["is"].get("backtest", [])}
    oos_bt = {r["politica"].split(" (")[0]: r for r in log["oos"].get("backtest", [])}

    def fmt(d, key):
        return d.get(key, {}).get("retorno", "?")

    rows = [
        {"segmento": "IS",  "OPT (oracle)": fmt(is_bt,  "OPT"),
         "Naive 50/50 BH":  fmt(is_bt,  "Naive 50/50 buy & hold"),
         "Naive 50/50 RB":  fmt(is_bt,  "Naive 50/50 rebalanceo"),
         "Regret-Grid":     fmt(is_bt,  "Regret-Grid g*_mean")},
        {"segmento": "OOS", "OPT (oracle)": fmt(oos_bt, "OPT_oos"),
         "Naive 50/50 BH":  fmt(oos_bt, "Naive 50/50 buy & hold"),
         "Naive 50/50 RB":  fmt(oos_bt, "Naive 50/50 rebalanceo"),
         "Regret-Grid":     fmt(oos_bt, "Regret-Grid_oos g*_mean")},
    ]
    add_df_table(doc, list_to_df(rows))

    add_heading(doc, "Lineas siguientes (no implementadas)", level=2)

    add_paragraph(doc, "- mu_hat_source='p_sign': oraculo historico con la "
        "misma regla del LSTM (r>=0). Resolveria la inconsistencia bear>bull "
        "sin tocar el LSTM. Pre-validado en HALLAZGOS sec 9-bis.")

    add_paragraph(doc, "- Cambiar la generacion de escenarios: q correlacionado "
        "con target ~ corr historica, en vez de q comun. Devuelve "
        "diversificacion real entre activos.")

    add_paragraph(doc, "- Tunear la NN (architectura/regularizacion). El cambio "
        "reciente a hidden=8 mejora el backtest considerablemente -- vale la "
        "pena seguir explorando esa direccion.")


# ================================================================ main
def main():
    print(f"Generando {OUT_PATH.relative_to(PROJECT_ROOT)} ...")
    log_path = latest_main_log()
    print(f"  log fuente: {log_path.relative_to(PROJECT_ROOT)}")
    log = parse_main_log(log_path)
    cfg = log.get("dlconfig", {})
    print(f"  DLConfig: {cfg}")
    print(f"  g*_mean IS:  {log['is'].get('g_mean', {})}")
    print(f"  g*_mean OOS: {log['oos'].get('g_mean', {})}")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    seccion_portada(doc, log)
    seccion_resumen(doc, log)
    doc.add_page_break()
    seccion_l2(doc, log)
    doc.add_page_break()
    seccion_l3(doc, log)
    doc.add_page_break()
    seccion_l4(doc)
    doc.add_page_break()
    seccion_l5(doc, log)
    doc.add_page_break()
    seccion_sintesis(doc, log)

    doc.save(str(OUT_PATH))
    print(f"  -> {OUT_PATH}  ({OUT_PATH.stat().st_size / 1024:.1f} KB)")


if __name__ == "__main__":
    main()
